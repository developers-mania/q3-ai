"""Sprint26 Q3 · corpus preparation.

Converts source documents (.docx, .pdf, .txt) into the plain text the pipeline
reads, then reports on whether the conversion kept the things the quarter depends
on.

WHY the pipeline only reads .txt:
    Session 01 is "the least work that produces something loadable" — one call to
    read_text(). Pointing it at a PDF imports an *extraction* problem into a
    session about retrieval. Extraction is named material for Sessions 02 and 03.
    Doing it here would add libraries to a deliberately tiny requirements.txt and
    would let a bad text layer depress the baseline for reasons that have nothing
    to do with chunking or indexing.

    So: convert once, in advance, as facilitator preparation. Commit the .txt
    checksums via tools/verify_corpus.py. The room never sees this script.

WHY the report matters more than the conversion:
    Naive PDF extraction silently destroys section numbering — running headers get
    interleaved, two-column layouts scramble reading order, and numbering held in a
    separate text run disappears. Session 03 recovers structure that Session 01
    discards at the chunking stage. If the numbering was never in the text, there
    is nothing to recover and the largest predicted jump in the quarter's number
    quietly does not happen. This script tells you before Wednesday, not in October.

Usage:
    pip install -r requirements-dev.txt
    python tools/prepare_corpus.py sources/dpa.pdf --name dpa-2019
    python tools/prepare_corpus.py sources/*.docx
    python tools/prepare_corpus.py --check-only          # re-report on corpus/*.txt
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import unicodedata
from collections import Counter
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
CORPUS = ROOT / "corpus"

# Strings the seed evaluation questions depend on being present. If a landmark is
# missing after conversion the extraction dropped or mangled it, and questions
# written against it will score MISS no matter how good retrieval gets.
#
# DPA s.43: statutes spell numbers out. "seventy-two hours" is the surface form,
# not "72 hours" — the same distinction the `expect` field exists for. Matched through
# flatten(), so hyphenation and the PDF's line wrapping do not produce false alarms.
LANDMARKS: dict[str, list[str]] = {
    "dpa-2019": [
        "seventy-two hours",     # s.43(1)(a) — controller to Commissioner
        "forty-eight hours",     # s.43(3)   — processor to controller
        "ninety days",           # s.56(5)   — complaint concluded, wrapped in the PDF
        "Data Commissioner",
    ],
    "ai-strategy-2025": [
        "pillars",
        "enablers",
    ],
}

SECTION_RE = re.compile(r"^\s*(\d{1,3})\.\s", re.MULTILINE)


def flatten(text: str) -> str:
    """Lowercase, collapse whitespace, treat hyphens as spaces.

    Deliberately the same rule as src.pipeline.normalise_for_match, restated rather
    than imported: tools/ must run before anything in src/ is trusted. A landmark that
    only fails because the PDF wrapped it mid-phrase is a false alarm, and a false
    alarm here costs a facilitator an evening looking for a better source.
    """
    return " ".join(text.replace("-", " ").split()).lower()


# --------------------------------------------------------------------------
# Conversion
# --------------------------------------------------------------------------

def from_docx(path: Path) -> str:
    """Paragraphs and table cells, in document order.

    WHY docx is the better source when you have the choice: it is structured XML,
    so numbering and headings come through as text rather than being inferred from
    glyph positions the way they are in a PDF.
    """
    try:
        import docx  # python-docx
    except ImportError:
        raise SystemExit("python-docx missing — pip install -r requirements-dev.txt")

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def from_pdf(path: Path) -> str:
    """pdftotext -layout if poppler is present, else pdfplumber.

    -layout preserves column structure. Without it, a two-column policy document
    interleaves the columns line by line and the output is unusable prose.
    """
    if shutil.which("pdftotext"):
        # WHY encoding= is explicit: text=True alone decodes with the locale codec
        # (cp1252 on Windows), which throws on the UTF-8 pdftotext was asked for.
        result = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            capture_output=True, text=True, encoding="utf-8", check=True,
        )
        return result.stdout

    try:
        import pdfplumber
    except ImportError:
        raise SystemExit("Neither pdftotext nor pdfplumber available.")

    with pdfplumber.open(str(path)) as pdf:
        return "\n".join((page.extract_text() or "") for page in pdf.pages)


def convert(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return from_docx(path)
    if suffix == ".pdf":
        return from_pdf(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise SystemExit(f"Unsupported source format: {path.name}")


def normalise(text: str) -> str:
    """UTF-8, Unix line endings, no smart punctuation, no runaway blank lines.

    WHY NFKC: PDFs carry ligatures (ﬁ, ﬂ) and non-breaking spaces that make an
    exact-match `expect` string fail for invisible reasons. Normalising once here
    means the room never debugs a MISS caused by a typographic space.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    for smart, plain in [("\u2018", "'"), ("\u2019", "'"), ("\u201c", '"'),
                         ("\u201d", '"'), ("\u2013", "-"), ("\u2014", "-"),
                         ("\u00a0", " ")]:
        text = text.replace(smart, plain)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + "\n"


# --------------------------------------------------------------------------
# The report — the part that actually earns this script
# --------------------------------------------------------------------------

def find_page_furniture(text: str, threshold: int = 5) -> list[tuple[str, int]]:
    """Short lines repeating many times: running headers, footers, page numbers.

    These are not fatal, but they pollute chunks and inflate TF-IDF weights for
    words that carry no meaning. Strip them before committing if the count is high.
    """
    lines = [ln.strip() for ln in text.split("\n")]
    candidates = Counter(ln for ln in lines if 0 < len(ln) < 60)
    return [(ln, n) for ln, n in candidates.most_common(8) if n >= threshold]


def report(name: str, text: str) -> bool:
    print(f"\n{'=' * 66}\n{name}.txt\n{'=' * 66}")
    print(f"  characters          {len(text):>10,}")
    print(f"  lines               {len(text.splitlines()):>10,}")

    sections = SECTION_RE.findall(text)
    numbers = sorted({int(s) for s in sections})
    print(f"  section markers     {len(sections):>10,}  "
          f"(distinct {len(numbers)}"
          + (f", range {numbers[0]}-{numbers[-1]}" if numbers else "")
          + ")")

    ok = True

    # Structure survival. This is the check that protects Session 03.
    if name.startswith("dpa") and len(numbers) < 40:
        print("\n  STRUCTURE WARNING — the Act has 75 sections; few markers survived.")
        print("  Session 03 recovers section boundaries from this text. If they are")
        print("  not here, that session has nothing to work with. Try a different")
        print("  source (a DOCX if you have one) before accepting this.")
        ok = False

    furniture = find_page_furniture(text)
    if furniture:
        print("\n  Repeated short lines — likely page furniture:")
        for line, count in furniture:
            print(f"    {count:>4}x  {line[:56]!r}")
        print("  Not fatal. Strip them if the counts are large.")

    flat = flatten(text)
    missing = [s for s in LANDMARKS.get(name, []) if flatten(s) not in flat]
    if LANDMARKS.get(name):
        found = len(LANDMARKS[name]) - len(missing)
        print(f"\n  landmarks           {found}/{len(LANDMARKS[name])} present")
        for s in missing:
            print(f"    MISSING  {s!r}")
        if missing:
            print("  A missing landmark means questions written against it score MISS")
            print("  regardless of retrieval quality. Check the source before Wednesday.")
            ok = False

    return ok


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Prepare corpus text from source documents")
    parser.add_argument("sources", nargs="*", type=Path, help=".docx, .pdf or .txt")
    parser.add_argument("--name", help="output stem (defaults to the source stem)")
    parser.add_argument("--check-only", action="store_true", help="re-report on corpus/*.txt")
    args = parser.parse_args(argv[1:])

    CORPUS.mkdir(exist_ok=True)
    all_ok = True

    if args.check_only:
        files = sorted(CORPUS.glob("*.txt"))
        if not files:
            print(f"No .txt files in {CORPUS}.")
            return 1
        for path in files:
            all_ok &= report(path.stem, path.read_text(encoding="utf-8"))
    else:
        if not args.sources:
            parser.error("give at least one source file, or use --check-only")
        if args.name and len(args.sources) > 1:
            parser.error("--name only makes sense with a single source")
        for source in args.sources:
            if not source.exists():
                print(f"Not found: {source}")
                return 1
            stem = args.name or source.stem
            text = normalise(convert(source))
            # WHY newline="\n": without it Windows writes CRLF, so the same PDF through
            # the same tool checksums differently per platform and the pin cries wolf.
            (CORPUS / f"{stem}.txt").write_text(text, encoding="utf-8", newline="\n")
            print(f"{source}  ->  corpus/{stem}.txt")
            all_ok &= report(stem, text)

    print(f"\n{'=' * 66}")
    if all_ok:
        print("Conversion looks sound. Next: python tools/verify_corpus.py")
    else:
        print("Fix the warnings above before pinning the corpus.")
    return 0 if all_ok else 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
